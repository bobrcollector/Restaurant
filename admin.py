import sys
import sqlite3
from PyQt6.QtWidgets import (QApplication, QMainWindow, QVBoxLayout, QWidget, QTableWidget, QTableWidgetItem,
                             QPushButton, QLineEdit, QLabel, QMessageBox, QDialog, QFormLayout, QComboBox, QTabWidget,
                             QCalendarWidget, QDialogButtonBox, QTimeEdit, QHeaderView)
from PyQt6.QtCore import Qt, QDate, QTime
from PyQt6.QtWidgets import QHBoxLayout
from PyQt6.QtGui import QFont
import docx


class AdminWindow(QMainWindow):
    def __init__(self, login_window):
        super().__init__()

        self.setWindowTitle("Администраторская панель")
        self.setGeometry(100, 100, 800, 600)

        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)

        self.layout = QVBoxLayout(self.central_widget)

        self.connection = sqlite3.connect('restaurant.db')
        self.cursor = self.connection.cursor()
        self.login_window = login_window

        self.init_ui()

    def init_ui(self):
        self.tab_widget = QTabWidget()
        self.layout.addWidget(self.tab_widget)

        self.create_users_tab()
        self.create_menu_tab()
        self.create_tables_tab()
        self.create_reservations_tab()
        self.create_shifts_tab()
        self.create_categories_tab()
        self.create_reports_tab()

        self.logout_button = QPushButton("Выход")
        self.logout_button.clicked.connect(self.logout)
        self.layout.addWidget(self.logout_button)

    def logout(self):
        self.close()
        self.login_window.show()

    def create_users_tab(self):
        users_tab = QWidget()
        users_layout = QVBoxLayout(users_tab)

        users_layout.addWidget(QLabel("Пользователи"))
        self.users_table = QTableWidget(0, 5)
        self.users_table.setHorizontalHeaderLabels(["ID", "Имя пользователя", "Пароль", "Email", "Роль"])
        self.users_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.users_table.setSelectionMode(QTableWidget.SelectionMode.SingleSelection)
        self.users_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        users_layout.addWidget(self.users_table)
        self.load_users_data()

        self.add_user_button = QPushButton("Добавить пользователя")
        self.add_user_button.clicked.connect(self.add_user)
        users_layout.addWidget(self.add_user_button)

        self.edit_user_button = QPushButton("Изменить пользователя")
        self.edit_user_button.clicked.connect(self.edit_user)
        users_layout.addWidget(self.edit_user_button)

        self.delete_user_button = QPushButton("Удалить пользователя")
        self.delete_user_button.clicked.connect(self.delete_user)
        users_layout.addWidget(self.delete_user_button)

        self.tab_widget.addTab(users_tab, "Пользователи")

    def create_menu_tab(self):
        menu_tab = QWidget()
        menu_layout = QVBoxLayout(menu_tab)

        menu_layout.addWidget(QLabel("Меню"))
        self.menu_table = QTableWidget(0, 4)
        self.menu_table.setHorizontalHeaderLabels(["ID", "Название блюда", "Цена", "Категория"])
        self.menu_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.menu_table.setSelectionMode(QTableWidget.SelectionMode.SingleSelection)
        self.menu_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        menu_layout.addWidget(self.menu_table)
        self.load_menu_data()

        self.add_dish_button = QPushButton("Добавить блюдо")
        self.add_dish_button.clicked.connect(self.add_dish)
        menu_layout.addWidget(self.add_dish_button)

        self.edit_dish_button = QPushButton("Изменить блюдо")
        self.edit_dish_button.clicked.connect(self.edit_dish)
        menu_layout.addWidget(self.edit_dish_button)

        self.delete_dish_button = QPushButton("Удалить блюдо")
        self.delete_dish_button.clicked.connect(self.delete_dish)
        menu_layout.addWidget(self.delete_dish_button)

        self.tab_widget.addTab(menu_tab, "Меню")

    def create_tables_tab(self):
        tables_tab = QWidget()
        tables_layout = QVBoxLayout(tables_tab)

        tables_layout.addWidget(QLabel("Столики"))
        self.tables_table = QTableWidget(0, 3)
        self.tables_table.setHorizontalHeaderLabels(["Номер столика", "Вместимость", "Статус"])
        self.tables_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.tables_table.setSelectionMode(QTableWidget.SelectionMode.SingleSelection)
        self.tables_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        tables_layout.addWidget(self.tables_table)
        self.load_tables_data()

        self.add_table_button = QPushButton("Добавить столик")
        self.add_table_button.clicked.connect(self.add_table)
        tables_layout.addWidget(self.add_table_button)

        self.edit_table_button = QPushButton("Изменить столик")
        self.edit_table_button.clicked.connect(self.edit_table)
        tables_layout.addWidget(self.edit_table_button)

        self.delete_table_button = QPushButton("Удалить столик")
        self.delete_table_button.clicked.connect(self.delete_table)
        tables_layout.addWidget(self.delete_table_button)

        self.tab_widget.addTab(tables_tab, "Столики")

    def create_reservations_tab(self):
        reservations_tab = QWidget()
        reservations_layout = QVBoxLayout(reservations_tab)

        reservations_layout.addWidget(QLabel("Бронирования"))
        self.reservations_table = QTableWidget(0, 5)
        self.reservations_table.setHorizontalHeaderLabels(["ID", "Номер столика", "Имя клиента", "Дата бронирования", "Статус"])
        self.reservations_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.reservations_table.setSelectionMode(QTableWidget.SelectionMode.SingleSelection)
        self.reservations_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        reservations_layout.addWidget(self.reservations_table)
        self.load_reservations_data()

        self.add_reservation_button = QPushButton("Добавить бронирование")
        self.add_reservation_button.clicked.connect(self.add_reservation)
        reservations_layout.addWidget(self.add_reservation_button)

        self.edit_reservation_button = QPushButton("Изменить бронирование")
        self.edit_reservation_button.clicked.connect(self.edit_reservation)
        reservations_layout.addWidget(self.edit_reservation_button)

        self.delete_reservation_button = QPushButton("Удалить бронирование")
        self.delete_reservation_button.clicked.connect(self.delete_reservation)
        reservations_layout.addWidget(self.delete_reservation_button)

        self.tab_widget.addTab(reservations_tab, "Бронирования")

    def create_shifts_tab(self):
        shifts_tab = QWidget()
        shifts_layout = QVBoxLayout(shifts_tab)

        shifts_layout.addWidget(QLabel("Смены"))
        self.shifts_table = QTableWidget(0, 5)
        self.shifts_table.setHorizontalHeaderLabels(["ID", "Пользователь", "Дата смены", "Время начала", "Время окончания"])
        self.shifts_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.shifts_table.setSelectionMode(QTableWidget.SelectionMode.SingleSelection)
        self.shifts_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        shifts_layout.addWidget(self.shifts_table)
        self.load_shifts_data()

        self.add_shift_button = QPushButton("Добавить смену")
        self.add_shift_button.clicked.connect(self.add_shift)
        shifts_layout.addWidget(self.add_shift_button)

        self.edit_shift_button = QPushButton("Изменить смену")
        self.edit_shift_button.clicked.connect(self.edit_shift)
        shifts_layout.addWidget(self.edit_shift_button)

        self.delete_shift_button = QPushButton("Удалить смену")
        self.delete_shift_button.clicked.connect(self.delete_shift)
        shifts_layout.addWidget(self.delete_shift_button)

        self.tab_widget.addTab(shifts_tab, "Смены")

    def create_categories_tab(self):
        categories_tab = QWidget()
        categories_layout = QVBoxLayout(categories_tab)

        categories_layout.addWidget(QLabel("Категории блюд"))
        self.categories_table = QTableWidget(0, 2)
        self.categories_table.setHorizontalHeaderLabels(["ID", "Название категории"])
        self.categories_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.categories_table.setSelectionMode(QTableWidget.SelectionMode.SingleSelection)
        self.categories_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        categories_layout.addWidget(self.categories_table)
        self.load_categories_data()

        self.add_category_button = QPushButton("Добавить категорию")
        self.add_category_button.clicked.connect(self.add_category)
        categories_layout.addWidget(self.add_category_button)

        self.edit_category_button = QPushButton("Изменить категорию")
        self.edit_category_button.clicked.connect(self.edit_category)
        categories_layout.addWidget(self.edit_category_button)

        self.delete_category_button = QPushButton("Удалить категорию")
        self.delete_category_button.clicked.connect(self.delete_category)
        categories_layout.addWidget(self.delete_category_button)

        self.tab_widget.addTab(categories_tab, "Категории блюд")

    def create_reports_tab(self):
        reports_tab = QWidget()
        reports_layout = QVBoxLayout(reports_tab)

        # Общая сумма
        self.total_sales_label = QLabel("Общая сумма: ")
        self.total_sales_label.setFont(QFont("Arial", 12))
        reports_layout.addWidget(self.total_sales_label)

        # Популярные блюда
        self.popular_dishes_label = QLabel("Популярные блюда:")
        self.popular_dishes_label.setFont(QFont("Arial", 12))
        reports_layout.addWidget(self.popular_dishes_label)

        self.popular_dishes_list = QLabel("")
        reports_layout.addWidget(self.popular_dishes_list)

        # У кого сколько отработано смен
        self.shifts_label = QLabel("Смены пользователей:")
        self.shifts_label.setFont(QFont("Arial", 12))
        reports_layout.addWidget(self.shifts_label)

        self.shifts_list = QLabel("")
        reports_layout.addWidget(self.shifts_list)

        # Кнопка для экспорта в Word
        export_button = QPushButton("Экспортировать отчет в Word")
        export_button.clicked.connect(self.export_report_to_word)
        reports_layout.addWidget(export_button)

        # Display the reports immediately
        self.show_sales_report()
        self.show_popular_dishes_report()
        self.show_shifts_report()

        self.tab_widget.addTab(reports_tab, "Отчеты")

    def show_sales_report(self):
        self.cursor.execute("SELECT SUM(total_amount) FROM Orders WHERE status = 'оплачен'")
        total_sales = self.cursor.fetchone()[0] or 0
        self.total_sales_label.setText(f"Общая сумма: {total_sales}")

    def show_popular_dishes_report(self):
        self.cursor.execute("""
            SELECT dish_name, SUM(quantity)
            FROM Orders JOIN Menu ON Orders.dish_id = Menu.dish_id
            WHERE Orders.status = 'оплачен'
            GROUP BY dish_name
            ORDER BY SUM(quantity) DESC
        """)
        popular_dishes = self.cursor.fetchall()
        report = "\n".join([f"{dish[0]}: {dish[1]} шт." for dish in popular_dishes])
        self.popular_dishes_list.setText(report)

    def show_shifts_report(self):
        self.cursor.execute("""
            SELECT Users.username, COUNT(Shifts.shift_id)
            FROM Shifts JOIN Users ON Shifts.user_id = Users.user_id
            GROUP BY Users.username
        """)
        shifts = self.cursor.fetchall()
        report = "\n".join([f"{user[0]}: {user[1]} смен" for user in shifts])
        self.shifts_list.setText(report)

    def export_report_to_word(self):
        doc = docx.Document()
        doc.add_heading('Отчет', level=1)

        # Общая сумма
        total_sales = self.total_sales_label.text()
        doc.add_paragraph(total_sales)

        # Популярные блюда
        doc.add_heading('Популярные блюда:', level=2)
        popular_dishes = self.popular_dishes_list.text()
        doc.add_paragraph(popular_dishes)

        # Смены пользователей
        doc.add_heading('Смены пользователей:', level=2)
        shifts = self.shifts_list.text()
        doc.add_paragraph(shifts)

        # Сохранение файла
        doc.save('Отчет.docx')
        QMessageBox.information(self, "Экспорт", "Отчет успешно экспортирован в Word.")

    def load_users_data(self):
        self.cursor.execute("SELECT * FROM Users")
        users = self.cursor.fetchall()
        self.users_table.setRowCount(len(users))
        for row, user in enumerate(users):
            for col, data in enumerate(user):
                self.users_table.setItem(row, col, QTableWidgetItem(str(data)))

    def load_menu_data(self):
        self.cursor.execute("SELECT Menu.dish_id, Menu.dish_name, Menu.price, Categories.category_name FROM Menu JOIN Categories ON Menu.category_id = Categories.category_id")
        menu = self.cursor.fetchall()
        self.menu_table.setRowCount(len(menu))
        for row, dish in enumerate(menu):
            for col, data in enumerate(dish):
                self.menu_table.setItem(row, col, QTableWidgetItem(str(data)))

    def load_tables_data(self):
        self.cursor.execute("SELECT table_number, capacity, status FROM Tables")
        tables = self.cursor.fetchall()
        self.tables_table.setRowCount(len(tables))
        for row, table in enumerate(tables):
            for col, data in enumerate(table):
                self.tables_table.setItem(row, col, QTableWidgetItem(str(data)))

    def load_reservations_data(self):
        self.cursor.execute("SELECT Reservations.reservation_id, Tables.table_number, Reservations.customer_name, Reservations.reservation_date, Reservations.status FROM Reservations JOIN Tables ON Reservations.table_id = Tables.table_id")
        reservations = self.cursor.fetchall()
        self.reservations_table.setRowCount(len(reservations))
        for row, reservation in enumerate(reservations):
            for col, data in enumerate(reservation):
                self.reservations_table.setItem(row, col, QTableWidgetItem(str(data)))

    def load_shifts_data(self):
        self.cursor.execute("SELECT Shifts.shift_id, Users.username, Shifts.shift_date, Shifts.start_time, Shifts.end_time FROM Shifts JOIN Users ON Shifts.user_id = Users.user_id")
        shifts = self.cursor.fetchall()
        self.shifts_table.setRowCount(len(shifts))
        for row, shift in enumerate(shifts):
            for col, data in enumerate(shift):
                self.shifts_table.setItem(row, col, QTableWidgetItem(str(data)))

    def load_categories_data(self):
        self.cursor.execute("SELECT * FROM Categories")
        categories = self.cursor.fetchall()
        self.categories_table.setRowCount(len(categories))
        for row, category in enumerate(categories):
            for col, data in enumerate(category):
                self.categories_table.setItem(row, col, QTableWidgetItem(str(data)))

    def add_user(self):
        dialog = AddUserDialog(self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            username, password, email, role = dialog.get_user_data()
            try:
                self.cursor.execute("INSERT INTO Users (username, password, email, role, created_at) VALUES (?, ?, ?, ?, ?)",
                                    (username, password, email, role, QDate.currentDate().toString(Qt.DateFormat.ISODate)))
                self.connection.commit()
                self.load_users_data()
            except sqlite3.Error as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось добавить пользователя: {e}")

    def edit_user(self):
        selected_row = self.users_table.currentRow()
        if selected_row < 0:
            QMessageBox.warning(self, "Предупреждение", "Пожалуйста, выберите пользователя для изменения.")
            return

        user_id = self.users_table.item(selected_row, 0).text()
        dialog = AddUserDialog(self, user_id)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            username, password, email, role = dialog.get_user_data()
            try:
                self.cursor.execute("UPDATE Users SET username = ?, password = ?, email = ?, role = ? WHERE user_id = ?",
                                    (username, password, email, role, user_id))
                self.connection.commit()
                self.load_users_data()
            except sqlite3.Error as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось изменить пользователя: {e}")

    def delete_user(self):
        selected_row = self.users_table.currentRow()
        if selected_row < 0:
            QMessageBox.warning(self, "Предупреждение", "Пожалуйста, выберите пользователя для удаления.")
            return

        user_id = self.users_table.item(selected_row, 0).text()
        reply = QMessageBox.question(self, "Подтверждение", "Вы точно хотите удалить пользователя?",
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            try:
                self.cursor.execute("DELETE FROM Users WHERE user_id = ?", (int(user_id),))
                self.connection.commit()
                self.load_users_data()
            except sqlite3.Error as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось удалить пользователя: {e}")

    def add_dish(self):
        dialog = AddDishDialog(self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            dish_name, price, category_id = dialog.get_dish_data()
            try:
                self.cursor.execute("SELECT 1 FROM Menu WHERE dish_name = ?", (dish_name,))
                if self.cursor.fetchone():
                    QMessageBox.warning(self, "Предупреждение", "Блюдо с таким названием уже существует.")
                    return
                self.cursor.execute("INSERT INTO Menu (dish_name, price, category_id) VALUES (?, ?, ?)",
                                    (dish_name, price, category_id))
                self.connection.commit()
                self.load_menu_data()
            except sqlite3.Error as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось добавить блюдо: {e}")

    def edit_dish(self):
        selected_row = self.menu_table.currentRow()
        if selected_row < 0:
            QMessageBox.warning(self, "Предупреждение", "Пожалуйста, выберите блюдо для изменения.")
            return

        dish_id = self.menu_table.item(selected_row, 0).text()
        dialog = AddDishDialog(self, dish_id)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            dish_name, price, category_id = dialog.get_dish_data()
            try:
                self.cursor.execute("UPDATE Menu SET dish_name = ?, price = ?, category_id = ? WHERE dish_id = ?",
                                    (dish_name, price, category_id, dish_id))
                self.connection.commit()
                self.load_menu_data()
            except sqlite3.Error as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось изменить блюдо: {e}")

    def delete_dish(self):
        selected_row = self.menu_table.currentRow()
        if selected_row < 0:
            QMessageBox.warning(self, "Предупреждение", "Пожалуйста, выберите блюдо для удаления.")
            return

        dish_id = self.menu_table.item(selected_row, 0).text()
        reply = QMessageBox.question(self, "Подтверждение", "Вы точно хотите удалить блюдо?",
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            try:
                self.cursor.execute("DELETE FROM Menu WHERE dish_id = ?", (int(dish_id),))
                self.connection.commit()
                self.load_menu_data()
            except sqlite3.Error as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось удалить блюдо: {e}")

    def add_table(self):
        dialog = AddTableDialog(self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            table_number, capacity, status = dialog.get_table_data()
            try:
                self.cursor.execute("SELECT 1 FROM Tables WHERE table_number = ?", (table_number,))
                if self.cursor.fetchone():
                    QMessageBox.warning(self, "Предупреждение", "Столик с таким номером уже существует.")
                    return
                self.cursor.execute("INSERT INTO Tables (table_number, capacity, status) VALUES (?, ?, ?)",
                                    (table_number, capacity, status))
                self.connection.commit()
                self.load_tables_data()
            except sqlite3.Error as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось добавить столик: {e}")

    def edit_table(self):
        selected_row = self.tables_table.currentRow()
        if selected_row < 0:
            QMessageBox.warning(self, "Предупреждение", "Пожалуйста, выберите столик для изменения.")
            return

        table_id = self.tables_table.item(selected_row, 0).text()
        dialog = AddTableDialog(self, table_id)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            table_number, capacity, status = dialog.get_table_data()
            try:
                self.cursor.execute("UPDATE Tables SET table_number = ?, capacity = ?, status = ? WHERE table_id = ?",
                                    (table_number, capacity, status, table_id))
                self.connection.commit()
                self.load_tables_data()
            except sqlite3.Error as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось изменить столик: {e}")

    def delete_table(self):
        selected_row = self.tables_table.currentRow()
        if selected_row < 0:
            QMessageBox.warning(self, "Предупреждение", "Пожалуйста, выберите столик для удаления.")
            return

        table_id = self.tables_table.item(selected_row, 0).text()
        reply = QMessageBox.question(self, "Подтверждение", "Вы точно хотите удалить столик?",
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            try:
                self.cursor.execute("DELETE FROM Tables WHERE table_id = ?", (int(table_id),))
                self.connection.commit()
                self.load_tables_data()
            except sqlite3.Error as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось удалить столик: {e}")

    def add_reservation(self):
        dialog = AddReservationDialog(self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            table_id, customer_name, reservation_date, status = dialog.get_reservation_data()
            try:
                self.cursor.execute("INSERT INTO Reservations (table_id, customer_name, reservation_date, status) VALUES (?, ?, ?, ?)",
                                    (table_id, customer_name, reservation_date, status))
                self.connection.commit()
                self.load_reservations_data()
            except sqlite3.Error as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось добавить бронирование: {e}")

    def edit_reservation(self):
        selected_row = self.reservations_table.currentRow()
        if selected_row < 0:
            QMessageBox.warning(self, "Предупреждение", "Пожалуйста, выберите бронирование для изменения.")
            return

        reservation_id = self.reservations_table.item(selected_row, 0).text()
        dialog = AddReservationDialog(self, reservation_id)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            table_id, customer_name, reservation_date, status = dialog.get_reservation_data()
            try:
                self.cursor.execute("UPDATE Reservations SET table_id = ?, customer_name = ?, reservation_date = ?, status = ? WHERE reservation_id = ?",
                                    (table_id, customer_name, reservation_date, status, reservation_id))
                self.connection.commit()
                self.load_reservations_data()
            except sqlite3.Error as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось изменить бронирование: {e}")

    def delete_reservation(self):
        selected_row = self.reservations_table.currentRow()
        if selected_row < 0:
            QMessageBox.warning(self, "Предупреждение", "Пожалуйста, выберите бронирование для удаления.")
            return

        reservation_id = self.reservations_table.item(selected_row, 0).text()
        reply = QMessageBox.question(self, "Подтверждение", "Вы точно хотите удалить бронирование?",
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            try:
                self.cursor.execute("DELETE FROM Reservations WHERE reservation_id = ?", (int(reservation_id),))
                self.connection.commit()
                self.load_reservations_data()
            except sqlite3.Error as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось удалить бронирование: {e}")

    def add_shift(self):
        dialog = AddShiftDialog(self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            user_id, shift_date, start_time, end_time = dialog.get_shift_data()
            try:
                self.cursor.execute("INSERT INTO Shifts (user_id, shift_date, start_time, end_time) VALUES (?, ?, ?, ?)",
                                    (user_id, shift_date, start_time, end_time))
                self.connection.commit()
                self.load_shifts_data()
            except sqlite3.Error as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось добавить смену: {e}")

    def edit_shift(self):
        selected_row = self.shifts_table.currentRow()
        if selected_row < 0:
            QMessageBox.warning(self, "Предупреждение", "Пожалуйста, выберите смену для изменения.")
            return

        shift_id = self.shifts_table.item(selected_row, 0).text()
        dialog = AddShiftDialog(self, shift_id)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            user_id, shift_date, start_time, end_time = dialog.get_shift_data()
            try:
                self.cursor.execute("UPDATE Shifts SET user_id = ?, shift_date = ?, start_time = ?, end_time = ? WHERE shift_id = ?",
                                    (user_id, shift_date, start_time, end_time, shift_id))
                self.connection.commit()
                self.load_shifts_data()
            except sqlite3.Error as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось изменить смену: {e}")

    def delete_shift(self):
        selected_row = self.shifts_table.currentRow()
        if selected_row < 0:
            QMessageBox.warning(self, "Предупреждение", "Пожалуйста, выберите смену для удаления.")
            return

        shift_id = self.shifts_table.item(selected_row, 0).text()
        reply = QMessageBox.question(self, "Подтверждение", "Вы точно хотите удалить смену?",
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            try:
                self.cursor.execute("DELETE FROM Shifts WHERE shift_id = ?", (int(shift_id),))
                self.connection.commit()
                self.load_shifts_data()
            except sqlite3.Error as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось удалить смену: {e}")

    def add_category(self):
        dialog = AddCategoryDialog(self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            category_name = dialog.get_category_data()
            try:
                self.cursor.execute("SELECT 1 FROM Categories WHERE category_name = ?", (category_name,))
                if self.cursor.fetchone():
                    QMessageBox.warning(self, "Предупреждение", "Категория с таким названием уже существует.")
                    return
                self.cursor.execute("INSERT INTO Categories (category_name) VALUES (?)", (category_name,))
                self.connection.commit()
                self.load_categories_data()
            except sqlite3.Error as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось добавить категорию: {e}")

    def edit_category(self):
        selected_row = self.categories_table.currentRow()
        if selected_row < 0:
            QMessageBox.warning(self, "Предупреждение", "Пожалуйста, выберите категорию для изменения.")
            return

        category_id = self.categories_table.item(selected_row, 0).text()
        dialog = AddCategoryDialog(self, category_id)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            category_name = dialog.get_category_data()
            try:
                self.cursor.execute("UPDATE Categories SET category_name = ? WHERE category_id = ?",
                                    (category_name, category_id))
                self.connection.commit()
                self.load_categories_data()
            except sqlite3.Error as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось изменить категорию: {e}")

    def delete_category(self):
        selected_row = self.categories_table.currentRow()
        if selected_row < 0:
            QMessageBox.warning(self, "Предупреждение", "Пожалуйста, выберите категорию для удаления.")
            return

        category_id = self.categories_table.item(selected_row, 0).text()
        reply = QMessageBox.question(self, "Подтверждение", "Вы точно хотите удалить категорию?",
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            try:
                self.cursor.execute("DELETE FROM Categories WHERE category_id = ?", (int(category_id),))
                self.connection.commit()
                self.load_categories_data()
            except sqlite3.Error as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось удалить категорию: {e}")

class AddUserDialog(QDialog):
    def __init__(self, parent=None, user_id=None):
        super().__init__(parent)
        self.setWindowTitle("Добавить пользователя")
        self.layout = QFormLayout(self)

        self.user_id = user_id
        self.username_input = QLineEdit(self)
        self.password_input = QLineEdit(self)
        self.email_input = QLineEdit(self)
        self.role_input = QComboBox(self)
        self.role_input.addItems(["Администратор", "Официант", "Повар"])

        self.layout.addRow("Имя пользователя", self.username_input)
        self.layout.addRow("Пароль", self.password_input)
        self.layout.addRow("Email", self.email_input)
        self.layout.addRow("Роль", self.role_input)

        self.buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel, self)
        self.buttons.accepted.connect(self.accept)
        self.buttons.rejected.connect(self.reject)
        self.layout.addWidget(self.buttons)

        if user_id:
            self.load_user_data()

    def load_user_data(self):
        connection = sqlite3.connect('restaurant.db')
        cursor = connection.cursor()
        cursor.execute("SELECT username, password, email, role FROM Users WHERE user_id = ?", (self.user_id,))
        user = cursor.fetchone()
        if user:
            self.username_input.setText(user[0])
            self.password_input.setText(user[1])
            self.email_input.setText(user[2])
            self.role_input.setCurrentText(user[3])
        connection.close()

    def get_user_data(self):
        return self.username_input.text(), self.password_input.text(), self.email_input.text(), self.role_input.currentText()

class AddDishDialog(QDialog):
    def __init__(self, parent=None, dish_id=None):
        super().__init__(parent)
        self.setWindowTitle("Добавить блюдо")
        self.layout = QFormLayout(self)

        self.dish_id = dish_id
        self.dish_name_input = QLineEdit(self)
        self.price_input = QLineEdit(self)
        self.category_input = QComboBox(self)

        # Загрузка категорий из базы данных
        self.cursor = parent.cursor
        self.cursor.execute("SELECT category_id, category_name FROM Categories")
        categories = self.cursor.fetchall()
        for category in categories:
            self.category_input.addItem(category[1], category[0])

        self.layout.addRow("Название блюда", self.dish_name_input)
        self.layout.addRow("Цена", self.price_input)
        self.layout.addRow("Категория", self.category_input)

        self.buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel, self)
        self.buttons.accepted.connect(self.accept)
        self.buttons.rejected.connect(self.reject)
        self.layout.addWidget(self.buttons)

        if dish_id:
            self.load_dish_data()

    def load_dish_data(self):
        connection = sqlite3.connect('restaurant.db')
        cursor = connection.cursor()
        cursor.execute("SELECT dish_name, price, category_id FROM Menu WHERE dish_id = ?", (self.dish_id,))
        dish = cursor.fetchone()
        if dish:
            self.dish_name_input.setText(dish[0])
            self.price_input.setText(str(dish[1]))
            self.category_input.setCurrentIndex(self.category_input.findData(dish[2]))
        connection.close()

    def get_dish_data(self):
        dish_name = self.dish_name_input.text().strip()
        price_text = self.price_input.text().strip()
        category_id = self.category_input.currentData()

        if not dish_name:
            QMessageBox.warning(self, "Предупреждение", "Пожалуйста, введите название блюда.")
            return None, None, None

        if not price_text:
            QMessageBox.warning(self, "Предупреждение", "Пожалуйста, введите цену.")
            return None, None, None

        try:
            price = float(price_text)
        except ValueError:
            QMessageBox.warning(self, "Предупреждение", "Пожалуйста, введите корректную цену.")
            return None, None, None

        return dish_name, price, category_id

class AddTableDialog(QDialog):
    def __init__(self, parent=None, table_id=None):
        super().__init__(parent)
        self.setWindowTitle("Добавить столик")
        self.layout = QFormLayout(self)

        self.table_id = table_id
        self.table_number_input = QLineEdit(self)
        self.capacity_input = QLineEdit(self)
        self.status_input = QComboBox(self)
        self.status_input.addItems(["свободен", "занят", "забронирован"])

        self.layout.addRow("Номер столика", self.table_number_input)
        self.layout.addRow("Вместимость", self.capacity_input)
        self.layout.addRow("Статус", self.status_input)

        self.buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel, self)
        self.buttons.accepted.connect(self.accept)
        self.buttons.rejected.connect(self.reject)
        self.layout.addWidget(self.buttons)

        if table_id:
            self.load_table_data()

    def load_table_data(self):
        connection = sqlite3.connect('restaurant.db')
        cursor = connection.cursor()
        cursor.execute("SELECT table_number, capacity, status FROM Tables WHERE table_id = ?", (self.table_id,))
        table = cursor.fetchone()
        if table:
            self.table_number_input.setText(str(table[0]))
            self.capacity_input.setText(str(table[1]))
            self.status_input.setCurrentText(table[2])
        connection.close()

    def get_table_data(self):
        return int(self.table_number_input.text()), int(self.capacity_input.text()), self.status_input.currentText()

class AddReservationDialog(QDialog):
    def __init__(self, parent=None, reservation_id=None):
        super().__init__(parent)
        self.setWindowTitle("Добавить бронирование")
        self.layout = QFormLayout(self)

        self.reservation_id = reservation_id
        self.table_id_input = QComboBox(self)
        self.cursor = parent.cursor
        self.cursor.execute("SELECT table_id, table_number FROM Tables")
        tables = self.cursor.fetchall()
        for table in tables:
            self.table_id_input.addItem(f"{table[1]}", table[0])

        self.customer_name_input = QLineEdit(self)
        self.reservation_date_input = QCalendarWidget(self)
        self.status_input = QComboBox(self)
        self.status_input.addItems(["активно", "отменено"])

        self.layout.addRow("Номер столика", self.table_id_input)
        self.layout.addRow("Имя клиента", self.customer_name_input)
        self.layout.addRow("Дата бронирования", self.reservation_date_input)
        self.layout.addRow("Статус", self.status_input)

        self.buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel, self)
        self.buttons.accepted.connect(self.accept)
        self.buttons.rejected.connect(self.reject)
        self.layout.addWidget(self.buttons)

        if reservation_id:
            self.load_reservation_data()

    def load_reservation_data(self):
        connection = sqlite3.connect('restaurant.db')
        cursor = connection.cursor()
        cursor.execute("SELECT table_id, customer_name, reservation_date, status FROM Reservations WHERE reservation_id = ?", (self.reservation_id,))
        reservation = cursor.fetchone()
        if reservation:
            self.table_id_input.setCurrentText(str(reservation[0]))
            self.customer_name_input.setText(reservation[1])
            self.reservation_date_input.setSelectedDate(QDate.fromString(reservation[2], Qt.DateFormat.ISODate))
            self.status_input.setCurrentText(reservation[3])
        connection.close()

    def get_reservation_data(self):
        return self.table_id_input.currentData(), self.customer_name_input.text(), self.reservation_date_input.selectedDate().toString(Qt.DateFormat.ISODate), self.status_input.currentText()

class AddShiftDialog(QDialog):
    def __init__(self, parent=None, shift_id=None):
        super().__init__(parent)
        self.setWindowTitle("Добавить смену")
        self.layout = QFormLayout(self)

        self.shift_id = shift_id
        self.user_id_input = QComboBox(self)
        self.cursor = parent.cursor
        self.cursor.execute("SELECT user_id, username FROM Users")
        users = self.cursor.fetchall()
        for user in users:
            self.user_id_input.addItem(f"{user[1]}", user[0])

        self.shift_date_input = QCalendarWidget(self)
        self.start_time_input = QTimeEdit(self)
        self.end_time_input = QTimeEdit(self)

        self.layout.addRow("Пользователь", self.user_id_input)
        self.layout.addRow("Дата смены", self.shift_date_input)
        self.layout.addRow("Время начала", self.start_time_input)
        self.layout.addRow("Время окончания", self.end_time_input)

        self.buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel, self)
        self.buttons.accepted.connect(self.accept)
        self.buttons.rejected.connect(self.reject)
        self.layout.addWidget(self.buttons)

        if shift_id:
            self.load_shift_data()

    def load_shift_data(self):
        connection = sqlite3.connect('restaurant.db')
        cursor = connection.cursor()
        cursor.execute("SELECT user_id, shift_date, start_time, end_time FROM Shifts WHERE shift_id = ?", (self.shift_id,))
        shift = cursor.fetchone()
        if shift:
            self.user_id_input.setCurrentIndex(self.user_id_input.findData(shift[0]))
            self.shift_date_input.setSelectedDate(QDate.fromString(shift[1], Qt.DateFormat.ISODate))
            self.start_time_input.setTime(QTime.fromString(shift[2], Qt.DateFormat.ISODate))
            self.end_time_input.setTime(QTime.fromString(shift[3], Qt.DateFormat.ISODate))
        connection.close()

    def get_shift_data(self):
        user_id = self.user_id_input.currentData()
        shift_date = self.shift_date_input.selectedDate().toString(Qt.DateFormat.ISODate)
        start_time = self.start_time_input.time().toString(Qt.DateFormat.ISODate)
        end_time = self.end_time_input.time().toString(Qt.DateFormat.ISODate)

        if not user_id:
            QMessageBox.warning(self, "Предупреждение", "Пожалуйста, выберите пользователя.")
            return None, None, None, None

        if not shift_date:
            QMessageBox.warning(self, "Предупреждение", "Пожалуйста, выберите дату смены.")
            return None, None, None, None

        if not start_time:
            QMessageBox.warning(self, "Предупреждение", "Пожалуйста, выберите время начала.")
            return None, None, None, None

        if not end_time:
            QMessageBox.warning(self, "Предупреждение", "Пожалуйста, выберите время окончания.")
            return None, None, None, None

        return user_id, shift_date, start_time, end_time

class AddCategoryDialog(QDialog):
    def __init__(self, parent=None, category_id=None):
        super().__init__(parent)
        self.setWindowTitle("Добавить категорию")
        self.layout = QFormLayout(self)

        self.category_id = category_id
        self.category_name_input = QLineEdit(self)

        self.layout.addRow("Название категории", self.category_name_input)

        self.buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel, self)
        self.buttons.accepted.connect(self.accept)
        self.buttons.rejected.connect(self.reject)
        self.layout.addWidget(self.buttons)

        if category_id:
            self.load_category_data()

    def load_category_data(self):
        connection = sqlite3.connect('restaurant.db')
        cursor = connection.cursor()
        cursor.execute("SELECT category_name FROM Categories WHERE category_id = ?", (self.category_id,))
        category = cursor.fetchone()
        if category:
            self.category_name_input.setText(category[0])
        connection.close()

    def get_category_data(self):
        return self.category_name_input.text()

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = AdminWindow()
    window.show()
    sys.exit(app.exec())
